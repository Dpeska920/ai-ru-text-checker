import base64
from io import BytesIO

import pytest
from docx import Document

from app.models.requests import FactChange
from app.services.diff_service import DiffService


@pytest.fixture
def diff_service():
    return DiffService()


def decode_docx(base64_content: str) -> Document:
    data = base64.b64decode(base64_content)
    return Document(BytesIO(data))


class TestDiffService:
    def test_generate_clean_doc(self, diff_service):
        original = "Hello world"
        corrected = "Hello beautiful world"

        clean_b64, _ = diff_service.generate(original, corrected)
        doc = decode_docx(clean_b64)

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Hello beautiful world" in text

    def test_generate_diff_doc(self, diff_service):
        original = "Hello world"
        corrected = "Hello beautiful world"

        _, diff_b64 = diff_service.generate(original, corrected)
        doc = decode_docx(diff_b64)

        assert len(doc.paragraphs) > 0

    def test_preserve_paragraphs(self, diff_service):
        original = "First paragraph.\n\nSecond paragraph."
        corrected = "First paragraph.\n\nSecond paragraph modified."

        clean_b64, _ = diff_service.generate(original, corrected)
        doc = decode_docx(clean_b64)

        paragraphs = [p.text for p in doc.paragraphs]
        assert len(paragraphs) == 3

    def test_empty_texts(self, diff_service):
        clean_b64, diff_b64 = diff_service.generate("", "")
        doc = decode_docx(clean_b64)
        assert doc is not None

    def test_same_text_no_diff(self, diff_service):
        text = "Same text without changes"
        clean_b64, diff_b64 = diff_service.generate(text, text)

        clean_doc = decode_docx(clean_b64)
        diff_doc = decode_docx(diff_b64)

        clean_text = "\n".join(p.text for p in clean_doc.paragraphs)
        diff_text = "\n".join(p.text for p in diff_doc.paragraphs)

        assert text in clean_text
        assert text in diff_text

    def test_fact_changes_provided(self, diff_service):
        original = "Глава Tesla Дональд Трамп объявил о новом продукте."
        corrected = "Глава Tesla Илон Маск объявил о новом продукте."

        fact_changes = [
            FactChange(
                original="Дональд Трамп",
                corrected="Илон Маск",
                context="Глава Tesla",
            )
        ]

        clean_b64, diff_b64 = diff_service.generate(original, corrected, fact_changes)

        clean_doc = decode_docx(clean_b64)
        decode_docx(diff_b64)

        clean_text = "\n".join(p.text for p in clean_doc.paragraphs)
        assert "Илон Маск" in clean_text
        assert "Дональд Трамп" not in clean_text

    def test_multiple_paragraphs_with_changes(self, diff_service):
        original = "First line.\nSecond line with error.\nThird line."
        corrected = "First line.\nSecond line corrected.\nThird line."

        clean_b64, diff_b64 = diff_service.generate(original, corrected)

        clean_doc = decode_docx(clean_b64)
        paragraphs = [p.text for p in clean_doc.paragraphs]

        assert len(paragraphs) == 3
        assert "corrected" in paragraphs[1]

    def test_cyrillic_text(self, diff_service):
        original = "Привет мир"
        corrected = "Привет прекрасный мир"

        clean_b64, diff_b64 = diff_service.generate(original, corrected)

        clean_doc = decode_docx(clean_b64)
        text = "\n".join(p.text for p in clean_doc.paragraphs)

        assert "Привет прекрасный мир" in text

    def test_deleted_text_only(self, diff_service):
        original = "Hello beautiful world"
        corrected = "Hello world"

        clean_b64, diff_b64 = diff_service.generate(original, corrected)
        clean_doc = decode_docx(clean_b64)

        text = "\n".join(p.text for p in clean_doc.paragraphs)
        assert "beautiful" not in text

    def test_added_paragraph(self, diff_service):
        original = "First.\nThird."
        corrected = "First.\nSecond.\nThird."

        clean_b64, _ = diff_service.generate(original, corrected)
        clean_doc = decode_docx(clean_b64)

        paragraphs = [p.text for p in clean_doc.paragraphs]
        assert len(paragraphs) == 3
        assert "Second." in paragraphs
