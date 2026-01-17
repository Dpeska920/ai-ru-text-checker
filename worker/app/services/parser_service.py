import base64
import binascii
from io import BytesIO

from docx import Document
from pypdf import PdfReader


class ParserError(Exception):
    pass


class EmptyFileError(ParserError):
    pass


class CorruptedFileError(ParserError):
    pass


class UnsupportedFormatError(ParserError):
    pass


class ParserService:
    def parse(self, file_content: str, file_type: str) -> str:
        """Parse document and extract text."""
        if not file_content:
            raise EmptyFileError("File content is empty")

        try:
            data = base64.b64decode(file_content)
        except binascii.Error as e:
            raise CorruptedFileError(f"Invalid base64 encoding: {e}")

        if not data:
            raise EmptyFileError("File is empty")

        buffer = BytesIO(data)

        match file_type.lower():
            case "docx":
                return self._parse_docx(buffer)
            case "pdf":
                return self._parse_pdf(buffer)
            case "txt" | "md":
                return self._parse_text(data)
            case "doc":
                raise UnsupportedFormatError(
                    "Legacy .doc format not supported. Please convert to .docx"
                )
            case _:
                raise UnsupportedFormatError(f"Unsupported file type: {file_type}")

    def _parse_docx(self, buffer: BytesIO) -> str:
        try:
            doc = Document(buffer)
        except Exception as e:
            raise CorruptedFileError(f"Cannot parse DOCX file: {e}")

        paragraphs = [p.text for p in doc.paragraphs]

        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                tables_text.append("\t".join(cells))

        all_text = "\n".join(paragraphs)
        if tables_text:
            all_text += "\n\n" + "\n".join(tables_text)

        text = all_text.strip()
        if not text:
            raise EmptyFileError("File is empty")

        return text

    def _parse_pdf(self, buffer: BytesIO) -> str:
        try:
            reader = PdfReader(buffer)
        except Exception as e:
            raise CorruptedFileError(f"Cannot parse PDF file: {e}")

        if len(reader.pages) == 0:
            raise EmptyFileError("PDF file has no pages")

        text_parts = []
        for page in reader.pages:
            try:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            except Exception:
                continue

        result = "\n".join(text_parts).strip()
        if not result:
            raise EmptyFileError("Could not extract text from PDF")

        return result

    def _parse_text(self, data: bytes) -> str:
        encodings = ["utf-8", "cp1251", "latin-1"]

        for encoding in encodings:
            try:
                text = data.decode(encoding).strip()
                if text:
                    return text
            except (UnicodeDecodeError, LookupError):
                continue

        raise CorruptedFileError("Cannot decode text file. Unsupported encoding")
