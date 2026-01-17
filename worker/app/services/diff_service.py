import base64
import difflib
import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor

from app.models.requests import FactChange


class DiffService:
    COLOR_ADDED = RGBColor(0x00, 0x50, 0x00)  # dark green on bright green highlight
    COLOR_DELETED = RGBColor(0x00, 0x00, 0x00)  # black on red highlight
    COLOR_FACT = RGBColor(0x00, 0x00, 0x00)  # black on yellow highlight

    def generate(
        self,
        original: str,
        corrected: str,
        fact_changes: list[FactChange] | None = None,
    ) -> tuple[str, str]:
        """Generate clean and diff documents. Returns base64 encoded docx files."""
        clean_doc = self._create_clean_doc(corrected)
        diff_doc = self._create_diff_doc(original, corrected, fact_changes)

        return self._doc_to_base64(clean_doc), self._doc_to_base64(diff_doc)

    def _create_clean_doc(self, text: str) -> Document:
        doc = Document()
        for paragraph in text.split("\n"):
            doc.add_paragraph(paragraph)
        return doc

    def _create_diff_doc(
        self,
        original: str,
        corrected: str,
        fact_changes: list[FactChange] | None = None,
    ) -> Document:
        doc = Document()

        fact_originals = set()
        fact_corrected = set()
        if fact_changes:
            for fc in fact_changes:
                fact_originals.add(fc.original.lower())
                fact_corrected.add(fc.corrected.lower())

        original_paragraphs = original.split("\n")
        corrected_paragraphs = corrected.split("\n")

        para_matcher = difflib.SequenceMatcher(
            None, original_paragraphs, corrected_paragraphs
        )

        for tag, i1, i2, j1, j2 in para_matcher.get_opcodes():
            match tag:
                case "equal":
                    for para_text in original_paragraphs[i1:i2]:
                        para = doc.add_paragraph()
                        para.add_run(para_text)
                case "delete":
                    for para_text in original_paragraphs[i1:i2]:
                        para = doc.add_paragraph()
                        self._add_deleted_paragraph(para, para_text)
                case "insert":
                    for para_text in corrected_paragraphs[j1:j2]:
                        para = doc.add_paragraph()
                        self._add_inserted_paragraph(para, para_text, fact_corrected)
                case "replace":
                    for idx in range(max(i2 - i1, j2 - j1)):
                        para = doc.add_paragraph()
                        orig_para = (
                            original_paragraphs[i1 + idx] if i1 + idx < i2 else ""
                        )
                        corr_para = (
                            corrected_paragraphs[j1 + idx] if j1 + idx < j2 else ""
                        )

                        if not orig_para:
                            self._add_inserted_paragraph(para, corr_para, fact_corrected)
                        elif not corr_para:
                            self._add_deleted_paragraph(para, orig_para)
                        else:
                            self._add_diff_paragraph(
                                para, orig_para, corr_para, fact_originals, fact_corrected
                            )

        return doc

    def _add_deleted_paragraph(self, para, text: str) -> None:
        run = para.add_run(text)
        run.font.color.rgb = self.COLOR_DELETED
        run.font.highlight_color = WD_COLOR_INDEX.RED

    def _add_inserted_paragraph(self, para, text: str, fact_corrected: set) -> None:
        words = self._tokenize(text)
        for word in words:
            if word.lower() in fact_corrected:
                run = para.add_run(word)
                run.font.color.rgb = self.COLOR_FACT
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                run = para.add_run(word)
                run.font.color.rgb = self.COLOR_ADDED
                run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

    def _add_diff_paragraph(
        self,
        para,
        original: str,
        corrected: str,
        fact_originals: set,
        fact_corrected: set,
    ) -> None:
        orig_words = self._tokenize_words(original)
        corr_words = self._tokenize_words(corrected)

        matcher = difflib.SequenceMatcher(None, orig_words, corr_words, autojunk=False)

        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            match tag:
                case "equal":
                    for word in orig_words[i1:i2]:
                        para.add_run(word)
                case "delete":
                    for word in orig_words[i1:i2]:
                        is_fact = word.strip().lower() in fact_originals
                        run = para.add_run(word)
                        if is_fact:
                            run.font.color.rgb = self.COLOR_FACT
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        else:
                            run.font.color.rgb = self.COLOR_DELETED
                            run.font.highlight_color = WD_COLOR_INDEX.RED
                case "insert":
                    for word in corr_words[j1:j2]:
                        is_fact = word.strip().lower() in fact_corrected
                        run = para.add_run(word)
                        if is_fact:
                            run.font.color.rgb = self.COLOR_FACT
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        else:
                            run.font.color.rgb = self.COLOR_ADDED
                            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                case "replace":
                    orig_chunk = orig_words[i1:i2]
                    corr_chunk = corr_words[j1:j2]

                    orig_text = "".join(orig_chunk)
                    corr_text = "".join(corr_chunk)

                    similarity = difflib.SequenceMatcher(None, orig_text, corr_text).ratio()

                    if similarity > 0.6:
                        self._add_char_diff(para, orig_text, corr_text, fact_originals, fact_corrected)
                    else:
                        is_fact_replacement = (
                            orig_text.strip().lower() in fact_originals
                            or corr_text.strip().lower() in fact_corrected
                        )

                        for word in orig_chunk:
                            run = para.add_run(word)
                            if is_fact_replacement:
                                run.font.color.rgb = self.COLOR_FACT
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                            else:
                                run.font.color.rgb = self.COLOR_DELETED
                                run.font.highlight_color = WD_COLOR_INDEX.RED

                        for word in corr_chunk:
                            run = para.add_run(word)
                            if is_fact_replacement:
                                run.font.color.rgb = self.COLOR_FACT
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                            else:
                                run.font.color.rgb = self.COLOR_ADDED
                                run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

    def _add_char_diff(
        self,
        para,
        original: str,
        corrected: str,
        fact_originals: set,
        fact_corrected: set,
    ) -> None:
        """Character-level diff for similar strings - shows precise changes."""
        matcher = difflib.SequenceMatcher(None, original, corrected, autojunk=False)

        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            match tag:
                case "equal":
                    para.add_run(original[i1:i2])
                case "delete":
                    chunk = original[i1:i2]
                    run = para.add_run(chunk)
                    run.font.color.rgb = self.COLOR_DELETED
                    run.font.highlight_color = WD_COLOR_INDEX.RED
                case "insert":
                    chunk = corrected[j1:j2]
                    run = para.add_run(chunk)
                    run.font.color.rgb = self.COLOR_ADDED
                    run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                case "replace":
                    del_chunk = original[i1:i2]
                    ins_chunk = corrected[j1:j2]

                    run = para.add_run(del_chunk)
                    run.font.color.rgb = self.COLOR_DELETED
                    run.font.highlight_color = WD_COLOR_INDEX.RED

                    run = para.add_run(ins_chunk)
                    run.font.color.rgb = self.COLOR_ADDED
                    run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

    def _tokenize(self, text: str) -> list[str]:
        """
        Split text into tokens preserving whitespace and punctuation as separate tokens.
        This allows for more accurate word-level diffing.
        """
        return re.findall(r"\S+|\s+", text)

    def _tokenize_words(self, text: str) -> list[str]:
        """
        Split text into words with their trailing whitespace.
        Better for word-level comparison - keeps words with their spaces.
        Example: "Hello world" -> ["Hello ", "world"]
        """
        tokens = []
        current = ""
        for char in text:
            if char.isspace():
                current += char
            else:
                if current and current[-1].isspace():
                    tokens.append(current)
                    current = char
                else:
                    current += char
        if current:
            tokens.append(current)
        return tokens if tokens else [text]

    def _doc_to_base64(self, doc: Document) -> str:
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return base64.b64encode(buffer.read()).decode("utf-8")
